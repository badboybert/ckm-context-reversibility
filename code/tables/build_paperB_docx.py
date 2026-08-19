#!/usr/bin/env python
"""Build Paper_B_MANUSCRIPT.docx from an assembled MANUSCRIPT.md (python-docx; node/pandoc absent).
Embeds main figures F1-F6 and
Extended Data figures ED1-ED6 above their legend line. Arial 11 pt, DOUBLE spacing, US Letter,
continuous line numbers and page numbers (Genome Medicine general guidance requires all three;
round 6 found the file at 1.5 spacing with no header/footer part at all, so no page number existed).

FLAT HEADINGS: the title, section
headings and subheadings share the BODY font (Arial) and colour (black) and are distinguished only
by bold + a light size step (14/12/11 pt vs 11 pt body). The heading outline level is kept for
navigation/TOC; only the run typography is forced.
Usage: python build_paperB_docx.py [input.md] [output.docx]"""
import os, re, sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_FONT = "Arial"
BODY_PT = 11
BLACK = RGBColor(0, 0, 0)          # identical to body text
HEAD_PT = {0: 14, 1: 12, 2: 11}    # title / section / subsection — size step only

HERE = os.path.dirname(os.path.abspath(__file__))
MD = os.path.abspath(sys.argv[1]) if len(sys.argv) > 1 else os.path.join(HERE, "MANUSCRIPT.md")
OUT = os.path.abspath(sys.argv[2]) if len(sys.argv) > 2 else os.path.join(HERE, "Paper_B_MANUSCRIPT.docx")
FIG = os.path.join(HERE, "figures")

def fig_path(kind, num):
    fn = ("Figure_ED%s.png" % num) if kind == "ED" else ("Figure_%s.png" % num)
    return os.path.join(FIG, fn)

# legend-line -> rendered image. Matches "**Figure N." and "**Extended Data Fig(ure) N."
LEGEND_RE = re.compile(r"^\*\*(Extended Data Fig(?:ure)?\.?|Figure)\s+(\d+)\b")
def legend_image(line):
    m = LEGEND_RE.match(line)
    if not m:
        return None
    kind = "ED" if m.group(1).startswith("Extended") else "F"
    p = fig_path(kind, m.group(2))
    return p if os.path.exists(p) else None

doc = Document()
n = doc.styles["Normal"]
n.font.name = BODY_FONT; n.font.size = Pt(BODY_PT); n.font.color.rgb = BLACK
pf = n.paragraph_format; pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE; pf.space_after = Pt(6)

sect = doc.sections[0]
sect.page_width, sect.page_height = Pt(612), Pt(792)  # US Letter
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sect, m, Pt(72))
ln = OxmlElement("w:lnNumType"); ln.set(qn("w:countBy"), "1"); ln.set(qn("w:restart"), "continuous")
sect._sectPr.append(ln)

# Page numbers. python-docx has no field API, so the PAGE field is built from its three XML parts.
# Without this the .docx carries no footer part at all — a fact invisible to every text-based check,
# because a missing page number is a missing OBJECT, not a wrong string.
def _add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.font.name = BODY_FONT; run.font.size = Pt(10); run.font.color.rgb = BLACK
    beg = OxmlElement("w:fldChar"); beg.set(qn("w:fldCharType"), "begin")
    ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve"); ins.text = " PAGE "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for el in (beg, ins, end):
        run._r.append(el)
_add_page_number(sect.footer.paragraphs[0])

INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
# Markdown backslash-escapes must be UNESCAPED in plain-text runs, or the literal backslash ships to
# the reader: a title-page footnote marker written "Tan\*" rendered as "Tan\*", and a JSON-escape
# residue "\"ns\"" rendered with visible backslashes. Applied only to non-emphasis segments so that
# "**bold**"/"*italic*" still parse.
_ESC = re.compile(r'\\([*_|"\[\]#`\\])')
def _unescape(s):
    return _ESC.sub(r'\1', s)
def strip_links(t):
    t = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", t)
    return t.replace("`", "")
def add_runs(par, text):
    for seg in INLINE.split(strip_links(text)):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            par.add_run(_unescape(seg[2:-2])).bold = True
        elif seg.startswith("*") and seg.endswith("*"):
            par.add_run(_unescape(seg[1:-1])).italic = True
        else:
            par.add_run(_unescape(seg))

def flat_heading(text, level, center=False):
    """Heading rendered in the BODY font + colour; outline level kept for navigation.
    Inline **bold**/*italic* markers are honoured, but every run is forced to Arial/black/bold
    with only a size step, so headings never diverge from body typography."""
    h = doc.add_heading("", level=level)
    for seg in INLINE.split(strip_links(text)):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            run = h.add_run(seg[2:-2])
        elif seg.startswith("*") and seg.endswith("*"):
            run = h.add_run(seg[1:-1]); run.italic = True
        else:
            run = h.add_run(seg)
        run.font.name = BODY_FONT
        run.font.size = Pt(HEAD_PT[level])
        run.font.color.rgb = BLACK
        run.bold = True
    if center:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return h

# Split a markdown table row on UNESCAPED pipes only, then unescape "\|" -> "|".
# Splitting on every "|" silently truncates any cell containing an escaped pipe: the STROBE-MR
# checklist's item 11 says "every pooled \|β\| < 0.029 SD", which the naive split tore into extra
# fields that were then dropped by the `j < len(brow)` guard below -- the shipped .docx ended that
# cell mid-sentence on a dangling backslash, losing ~200 characters, while the .md was complete.
_CELL_SPLIT = re.compile(r"(?<!\\)\|")
def _split_row(r):
    return [c.strip().replace("\\|", "|") for c in _CELL_SPLIT.split(r.strip().strip("|"))]

def add_table(rows):
    cells = [_split_row(r) for r in rows]
    header, body = cells[0], cells[2:]
    ncol = len(header)
    for k, brow in enumerate(cells[2:]):
        if len(brow) > ncol:
            raise AssertionError(
                f"table row {k} has {len(brow)} fields vs {ncol} header columns — a cell would be "
                f"silently truncated: {brow!r}")
    # Table Grid = borders only, no cell shading (Genome Medicine's general table guidance forbids
    # colour/shading; the prior "Light Grid Accent 1" applied a body-cell fill). Round 7 P0-2.
    t = doc.add_table(rows=1, cols=len(header)); t.style = "Table Grid"
    for j, h in enumerate(header):
        p = t.rows[0].cells[j].paragraphs[0]; add_runs(p, h)
        for run in p.runs:
            run.bold = True
    for brow in body:
        rc = t.add_row().cells
        for j in range(len(header)):
            add_runs(rc[j].paragraphs[0], brow[j] if j < len(brow) else "")
    for row in t.rows:
        for c in row.cells:
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(1)
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

raw = open(MD, encoding="utf-8").read()
raw = re.sub(r"<!--.*?-->", "", raw, flags=re.DOTALL)
lines = raw.splitlines()
i = 0
while i < len(lines):
    ln_ = lines[i]
    if re.match(r"^\|.*\|", ln_):
        blk = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            blk.append(lines[i]); i += 1
        if len(blk) >= 2:
            add_table(blk)
        continue
    if ln_.startswith("# "):
        flat_heading(ln_[2:], 0, center=True)
    elif ln_.startswith("## "):
        flat_heading(ln_[3:], 1)
    elif ln_.startswith("### "):
        flat_heading(ln_[4:], 2)
    elif ln_.startswith("---"):
        pass
    elif re.match(r"^[-*] ", ln_):
        add_runs(doc.add_paragraph(style="List Bullet"), ln_[2:])
    elif re.match(r"^\d+\. ", ln_):
        # KEEP THE LITERAL NUMERAL. Word's "List Number" style strips "N. " from the paragraph text
        # and re-derives the number from the list definition, so the numbers exist only as
        # formatting: any text extraction loses them, the portal's PDF converter re-derives them,
        # and any earlier numbered list in the same document silently shifts the whole run. That is
        # exactly how a sibling manuscript shipped a reference list printed 7.-53. while every
        # in-text citation said [1]-[47] — invisible to every text-based check, because the
        # paragraphs contained no numerals at all. Render as a plain hanging-indent paragraph.
        _p = doc.add_paragraph()
        _p.paragraph_format.left_indent = Pt(24)
        _p.paragraph_format.first_line_indent = Pt(-24)
        add_runs(_p, ln_)
    elif ln_.strip() == "":
        pass
    else:
        img = legend_image(ln_)
        if img:
            pic = doc.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic.add_run().add_picture(img, width=Inches(6.5))
        add_runs(doc.add_paragraph(), ln_)
    i += 1

# core-properties author metadata: creator/author = Bertrand Chin-Ming Tan;
# clear the python-docx default description; fixed modified/created date (no datetime.now()).
_STAMP = datetime(2026, 7, 9)
cp = doc.core_properties
cp.author = "Bertrand Chin-Ming Tan"
cp.last_modified_by = "Bertrand Chin-Ming Tan"
cp.comments = ""            # clears default "generated by python-docx" description (dc:description)
cp.created = _STAMP
cp.modified = _STAMP

doc.save(OUT)
print("wrote", OUT, "(", len(doc.paragraphs), "paragraphs )")
